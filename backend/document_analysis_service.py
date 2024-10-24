import logging
import sqlite3
import uuid
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import List, Dict, Any, Optional
import PyPDF2
from docx import Document
import pandas as pd
import io
import numpy as np
import math
import io
import PyPDF2
from datetime import datetime


def parse_pdf_date(date_string):
    if not date_string:
        return None

    date_formats = [
        'D:%Y%m%d%H%M%S%z',  # Standard PDF date format
        '%Y-%m-%d %H:%M:%S',  # ISO format
        '%m/%d/%Y %H:%M:%S',  # US format
        '%d/%m/%Y %H:%M:%S',  # UK format
    ]

    for date_format in date_formats:
        try:
            return datetime.strptime(date_string.replace("'", ""), date_format)
        except ValueError:
            continue

    # If no format matches, return the original string
    return date_string


def analyze_pdf(file_content: bytes) -> tuple[str, dict]:
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    content = ""
    for page in pdf_reader.pages:
        content += page.extract_text() + "\n"

    metadata = pdf_reader.metadata if pdf_reader.metadata else {}

    parsed_metadata = {
        "num_rows": 0,
        "num_columns": 0,
        "column_names": [],
        "num_pages": len(pdf_reader.pages),
        "author": metadata.get('/Author'),
        "creator": metadata.get('/Creator'),
        "producer": metadata.get('/Producer'),
        "subject": metadata.get('/Subject'),
        "title": metadata.get('/Title'),
        "creation_date": None,
        "modification_date": None,
    }

    creation_date = parse_pdf_date(metadata.get('/CreationDate'))
    if isinstance(creation_date, datetime):
        parsed_metadata["creation_date"] = creation_date.strftime('%Y-%m-%d %H:%M:%S')
    else:
        parsed_metadata["creation_date"] = creation_date

    modification_date = parse_pdf_date(metadata.get('/ModDate'))
    if isinstance(modification_date, datetime):
        parsed_metadata["modification_date"] = modification_date.strftime('%Y-%m-%d %H:%M:%S')
    else:
        parsed_metadata["modification_date"] = modification_date

    return content, parsed_metadata

def analyze_word(file_content: bytes) -> tuple[str, dict]:
    doc = Document(io.BytesIO(file_content))
    content = "\n".join([para.text for para in doc.paragraphs])
    metadata = {
        "num_rows": 0,
        "num_columns": 0,
        "column_names": [],
        "num_paragraphs": len(doc.paragraphs),
        "num_tables": len(doc.tables),
    }
    return content, metadata


def safe_convert(value):
    if pd.isna(value):
        return None
    if isinstance(value, (int, float)):
        if math.isinf(value) or math.isnan(value):
            return str(value)
        return value
    return str(value)


def analyze_spreadsheet(file_content: bytes, file_extension: str) -> tuple[str, dict, List[List[Any]]]:
    if file_extension in ['xls', 'xlsx']:
        df = pd.read_excel(io.BytesIO(file_content))
    else:  # csv
        df = pd.read_csv(io.BytesIO(file_content))

    content = df.to_string()
    metadata = {
        "num_rows": len(df),
        "num_columns": len(df.columns),
        "column_names": df.columns.tolist(),
    }

    # Convert DataFrame to a list of lists for JSON serialization
    excel_data = [df.columns.tolist()] + df.values.tolist()

    # Safely convert all values
    excel_data = [[safe_convert(cell) for cell in row] for row in excel_data]

    return content, metadata, excel_data
